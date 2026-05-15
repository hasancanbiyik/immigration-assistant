"""
Translation Router
==================
Handles text and uploaded document translation with USCIS certification.

Translation priority:
  1. Gemini 2.5 Flash (high quality, confidence scoring) — if API key is set
  2. Helsinki-NLP OPUS-MT (local model, no API key needed) — fallback
"""

import os
import time
import logging
from io import BytesIO

from fastapi import APIRouter, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import StreamingResponse

from app.models.schemas import (
    TranslationRequest,
    TranslationResponse,
    DocumentTranslationResponse,
    SupportedLanguage,
)
from app.utils.document_parser import (
    DocumentParser,
    SUPPORTED_DOCUMENT_EXTENSIONS,
    is_supported_document,
)
from app.services.translation import OPUS_MT_MODELS

logger = logging.getLogger(__name__)
router = APIRouter()
document_parser = DocumentParser()


def _extract_pages_for_translation(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract pages from a document for translation.

    For PDFs: uses PyMuPDF's simple text extraction (page.get_text("text"))
    which preserves natural reading order better than block-sorted extraction.

    For DOCX/TXT: falls back to the standard DocumentParser which handles
    these formats well.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for i in range(len(doc)):
                text = doc[i].get_text("text").strip()
                if text:
                    pages.append({"page_number": i + 1, "text": text})
            return pages
        except Exception as e:
            logger.warning(f"fitz direct extraction failed ({e}), falling back to DocumentParser")

    # Fallback: use DocumentParser (good for DOCX, TXT, etc.)
    parsed = document_parser.parse_document(file_bytes, filename)
    return [
        {"page_number": p.page_number, "text": p.text}
        for p in parsed.pages
        if p.text.strip()
    ]


# ─── Text translation ─────────────────────────────────────────────────

@router.post("/text", response_model=TranslationResponse)
async def translate_text(request: Request, body: TranslationRequest):
    """
    Translate text between supported language pairs.

    Uses Gemini 2.5 Flash if available (better quality + confidence score),
    otherwise falls back to OPUS-MT local model.
    """
    if body.source_lang == body.target_lang:
        raise HTTPException(
            status_code=400,
            detail="Source and target languages must be different.",
        )

    start_time = time.time()
    gemini = request.app.state.gemini
    translation_service = request.app.state.translation

    confidence = None
    model_used = None

    # Try Gemini first
    gemini_result = await gemini.translate_text(
        text=body.text,
        source_lang=body.source_lang.value,
        target_lang=body.target_lang.value,
    )

    if gemini_result:
        translated_text = gemini_result["translated_text"]
        confidence = gemini_result["confidence"]
        model_used = gemini_result["model_used"]
    else:
        # Fall back to OPUS-MT
        try:
            opus_result = translation_service.translate_text(
                text=body.text,
                source_lang=body.source_lang.value,
                target_lang=body.target_lang.value,
                generate_certification=False,
            )
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"OPUS-MT translation failed: {e}")
            raise HTTPException(status_code=500, detail="Translation failed. Please try again.")
        translated_text = opus_result["translated_text"]
        model_used = opus_result["model_used"]

    processing_time = (time.time() - start_time) * 1000
    word_count = len(body.text.split())

    certification = None
    if body.generate_certification:
        certification = translation_service._generate_certification(
            source_lang=body.source_lang.value,
            target_lang=body.target_lang.value,
            word_count=word_count,
            model_id=model_used,
        )

    return TranslationResponse(
        original_text=body.text,
        translated_text=translated_text,
        source_lang=body.source_lang.value,
        target_lang=body.target_lang.value,
        model_used=model_used,
        certification_statement=certification,
        word_count=word_count,
        processing_time_ms=round(processing_time, 2),
        confidence=confidence,
    )


# ─── Document translation ─────────────────────────────────────────────

@router.post("/document", response_model=DocumentTranslationResponse)
async def translate_document(
    request: Request,
    file: UploadFile = File(...),
    source_lang: SupportedLanguage = Form(...),
    target_lang: SupportedLanguage = Form(default=SupportedLanguage.ENGLISH),
    generate_certification: bool = Form(default=True),
):
    """
    Translate an entire supported document page by page.

    Uses Gemini 2.5 Flash if available, otherwise OPUS-MT.
    Returns translated text per page + average confidence score.
    """
    if source_lang == target_lang:
        raise HTTPException(
            status_code=400,
            detail="Source and target languages must be different.",
        )

    if not is_supported_document(file.filename):
        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Supported file types: {supported}.",
        )

    start_time = time.time()
    file_bytes = await file.read()

    try:
        pages = _extract_pages_for_translation(file_bytes, file.filename)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to parse document: {str(e)}")

    if not pages:
        raise HTTPException(status_code=422, detail="No text could be extracted from this document.")

    gemini = request.app.state.gemini
    translation_service = request.app.state.translation

    translated_pages = []
    confidence_scores = []

    for page in pages:
        # Try Gemini
        result = await gemini.translate_text(
            text=page["text"],
            source_lang=source_lang.value,
            target_lang=target_lang.value,
        )
        if result:
            translated_pages.append({
                "page_number": page["page_number"],
                "original_text": page["text"],
                "translated_text": result["translated_text"],
            })
            if result.get("confidence") is not None:
                confidence_scores.append(result["confidence"])
        else:
            # OPUS-MT fallback per page
            opus_result = translation_service.translate_text(
                text=page["text"],
                source_lang=source_lang.value,
                target_lang=target_lang.value,
                generate_certification=False,
            )
            translated_pages.append({
                "page_number": page["page_number"],
                "original_text": page["text"],
                "translated_text": opus_result["translated_text"],
            })

    avg_confidence = (
        round(sum(confidence_scores) / len(confidence_scores), 2)
        if confidence_scores else None
    )

    model_id = (
        "gemini-2.5-flash"
        if gemini.is_available
        else OPUS_MT_MODELS.get((source_lang.value, target_lang.value), "opus-mt")
    )

    certification = None
    if generate_certification:
        total_words = sum(len(p["text"].split()) for p in pages)
        certification = translation_service._generate_certification(
            source_lang=source_lang.value,
            target_lang=target_lang.value,
            word_count=total_words,
            model_id=model_id,
            document_description=file.filename,
        )

    processing_time = (time.time() - start_time) * 1000

    return DocumentTranslationResponse(
        original_filename=file.filename,
        translated_pages=translated_pages,
        total_pages=len(translated_pages),
        certification_statement=certification,
        processing_time_ms=round(processing_time, 2),
        confidence=avg_confidence,
    )


# ─── Export (download as .docx or .pdf) ──────────────────────────────

@router.post("/export")
async def export_translation(
    translated_text: str = Form(...),
    certification: str = Form(default=""),
    original_filename: str = Form(default="translation"),
    fmt: str = Form(default="docx"),  # "docx" or "pdf"
):
    """
    Export translated text (+ optional certification) as .docx or .pdf.
    """
    fmt = fmt.lower().strip()
    if fmt not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'.")

    base_name = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename

    if fmt == "docx":
        content, media_type, ext = _build_docx(translated_text, certification, base_name)
    else:
        content, media_type, ext = _build_pdf(translated_text, certification, base_name)

    filename = f"{base_name}_translation.{ext}"
    return StreamingResponse(
        BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_docx(translated_text: str, certification: str, title: str) -> tuple:
    """Build a .docx file from translated text and optional certification."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(f"Translation: {title}", 0)
    heading.runs[0].font.color.rgb = RGBColor(15, 110, 86)  # brand green

    # Warning notice
    notice = doc.add_paragraph()
    notice_run = notice.add_run(
        "⚠ This is an AI-assisted translation draft. "
        "It must be reviewed and certified by a qualified human translator "
        "before submission to USCIS or any government authority."
    )
    notice_run.font.size = Pt(9)
    notice_run.font.color.rgb = RGBColor(100, 56, 6)
    notice.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # spacer

    # Translation body
    body_heading = doc.add_heading("Translation", level=1)
    for para_text in translated_text.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # Certification (if present)
    if certification and certification.strip():
        doc.add_page_break()
        doc.add_heading("USCIS Certification Statement", level=1)
        for line in certification.split("\n"):
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"


def _build_pdf(translated_text: str, certification: str, title: str) -> tuple:
    """
    Build a PDF file from translated text using fpdf2.

    Font strategy:
      - If DejaVuSans.ttf is available (installed via `fonts-dejavu-core`
        in the Dockerfile), use it. DejaVu covers Turkish, Spanish, Arabic,
        Greek, Cyrillic — i.e. every language pair we support except Chinese.
      - Otherwise fall back to fpdf2's built-in Helvetica (Latin-1 only) and
        replace unsupported characters with '?'. This keeps local dev working
        on machines that don't have DejaVu installed.

    For Chinese (`zh`), neither core Helvetica nor DejaVuSans covers CJK
    glyphs — those characters will still render as boxes/question marks.
    Full CJK support would require shipping a Noto CJK font (~10–15 MB);
    deferred until there's actual demand.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF export requires the fpdf2 library. Run: pip install fpdf2",
        )

    # ── Detect Unicode font ──────────────────────────────────────────
    _DEJAVU_DIR = "/usr/share/fonts/truetype/dejavu"
    _candidates = {
        "regular":      os.path.join(_DEJAVU_DIR, "DejaVuSans.ttf"),
        "bold":         os.path.join(_DEJAVU_DIR, "DejaVuSans-Bold.ttf"),
        "italic":       os.path.join(_DEJAVU_DIR, "DejaVuSans-Oblique.ttf"),
    }
    has_unicode_font = all(os.path.exists(p) for p in _candidates.values())

    pdf = FPDF()
    # set_margins must be called BEFORE add_page so epw is computed correctly
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    w = pdf.epw

    if has_unicode_font:
        pdf.add_font("DejaVu", "",  _candidates["regular"])
        pdf.add_font("DejaVu", "B", _candidates["bold"])
        pdf.add_font("DejaVu", "I", _candidates["italic"])
        body_font = "DejaVu"

        def render(text: str) -> str:
            return text  # fpdf2 with TTF handles Unicode natively
    else:
        body_font = "Helvetica"

        def render(text: str) -> str:
            """Strip chars Helvetica can't render. Only invoked in fallback path."""
            return text.encode("latin-1", errors="replace").decode("latin-1")

    # ── Title ────────────────────────────────────────────────────────
    pdf.set_font(body_font, "B", 16)
    pdf.set_text_color(15, 110, 86)
    pdf.multi_cell(w, 8, render(f"Translation: {title}"))
    pdf.ln(4)

    # ── Warning notice ───────────────────────────────────────────────
    pdf.set_font(body_font, "I", 9)
    pdf.set_text_color(150, 80, 0)
    notice = (
        "WARNING: AI-assisted translation draft. Must be reviewed and certified "
        "by a qualified human translator before USCIS submission."
    )
    pdf.multi_cell(w, 5, render(notice))
    pdf.ln(8)

    # ── Translation body ─────────────────────────────────────────────
    pdf.set_font(body_font, "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.multi_cell(w, 8, render("Translation"))
    pdf.set_font(body_font, "", 11)

    for para in translated_text.split("\n\n"):
        para = para.strip()
        if para:
            pdf.multi_cell(w, 6, render(para))
            pdf.ln(3)

    # ── Certification ────────────────────────────────────────────────
    if certification and certification.strip():
        pdf.add_page()
        pdf.set_font(body_font, "B", 12)
        pdf.multi_cell(w, 8, render("USCIS Certification Statement"))
        pdf.set_font(body_font, "", 10)
        for line in certification.split("\n"):
            pdf.multi_cell(w, 5, render(line))

    content = bytes(pdf.output())
    return content, "application/pdf", "pdf"


# ─── Languages ────────────────────────────────────────────────────────

@router.get("/languages")
async def get_supported_languages(request: Request):
    """Get all supported language pairs and their models."""
    translation_service = request.app.state.translation
    return {
        "supported_pairs": translation_service.get_supported_languages(),
        "loaded_models": translation_service.get_loaded_models(),
        "languages": {
            "tr": "Turkish", "es": "Spanish",
            "zh": "Chinese", "ar": "Arabic", "en": "English",
        },
    }
