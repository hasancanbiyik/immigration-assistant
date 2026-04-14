"""
Document Parsing Pipeline
=========================
Routes supported uploads by file extension and reuses the PDF parser's
chunking and metadata extraction for .pdf, .txt, and .docx files.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import hashlib

from docx import Document as DocxDocument

from app.utils.pdf_parser import PDFParser, ParsedDocument, ParsedPage


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".docx"}


class DocumentParser(PDFParser):
    """Parses supported upload formats into a shared ParsedDocument shape."""

    def parse_document(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        extension = self._get_extension(filename)

        if extension == ".pdf":
            return self.parse_pdf_bytes(file_bytes, filename)
        if extension == ".txt":
            return self._parse_text_document(file_bytes, filename)
        if extension == ".docx":
            return self._parse_docx_document(file_bytes, filename)

        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
        raise ValueError(
            f"Unsupported file type: {extension or 'unknown'}. "
            f"Supported types: {supported}."
        )

    def _parse_text_document(
        self, file_bytes: bytes, filename: str
    ) -> ParsedDocument:
        text = self._decode_text_bytes(file_bytes)
        return self._build_text_document(text, filename)

    def _parse_docx_document(
        self, file_bytes: bytes, filename: str
    ) -> ParsedDocument:
        doc = DocxDocument(BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        text = "\n\n".join(paragraphs + table_rows)
        return self._build_text_document(text, filename)

    def _build_text_document(self, text: str, filename: str) -> ParsedDocument:
        clean_text = self._clean_text(text)
        pages = [
            ParsedPage(
                page_number=1,
                text=clean_text,
                word_count=len(clean_text.split()),
            )
        ] if clean_text else []
        full_text = clean_text
        doc_id = hashlib.sha256(
            f"{filename}:{full_text[:500]}".encode()
        ).hexdigest()[:12]
        metadata = self._extract_metadata(full_text)
        chunks = self._create_chunks(pages, filename)

        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            total_pages=len(pages),
            pages=pages,
            chunks=chunks,
            metadata=metadata,
            full_text=full_text,
        )

    def _decode_text_bytes(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Failed to decode text file.")

    def _get_extension(self, filename: str) -> str:
        return Path(filename).suffix.lower()


def is_supported_document(filename: str | None) -> bool:
    if not filename:
        return False
    return Path(filename).suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS
